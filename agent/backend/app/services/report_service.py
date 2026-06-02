import os
import uuid
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from jinja2 import Template
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable

from app.core.config import settings


def _ensure_storage():
    Path(settings.storage_path, "images").mkdir(parents=True, exist_ok=True)
    Path(settings.storage_path, "reports").mkdir(parents=True, exist_ok=True)


def _styles():
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "CustomTitle",
            parent=base["Title"],
            fontSize=20,
            spaceAfter=12,
            textColor=colors.HexColor("#1a3c5e"),
        ),
        "heading": ParagraphStyle(
            "CustomHeading",
            parent=base["Heading2"],
            fontSize=14,
            spaceBefore=16,
            spaceAfter=8,
            textColor=colors.HexColor("#2c6496"),
        ),
        "body": ParagraphStyle(
            "CustomBody",
            parent=base["Normal"],
            fontSize=11,
            leading=16,
        ),
        "meta": ParagraphStyle(
            "Meta",
            parent=base["Normal"],
            fontSize=9,
            textColor=colors.grey,
        ),
    }


def _html_template() -> Template:
    return Template("""<!DOCTYPE html>
<html lang="vi">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{{ title }}</title>
<style>
  * { box-sizing: border-box; margin: 0; padding: 0; }
  body { font-family: 'Segoe UI', Arial, sans-serif; font-size: 14px; line-height: 1.6;
         color: #333; background: #f5f5f5; padding: 40px 20px; }
  .container { max-width: 900px; margin: 0 auto; background: #fff;
               border-radius: 8px; box-shadow: 0 2px 12px rgba(0,0,0,0.1); overflow: hidden; }
  .header { background: linear-gradient(135deg, #1a3c5e, #2c6496); color: #fff; padding: 32px 40px; }
  .header h1 { font-size: 24px; margin-bottom: 8px; }
  .header .meta { opacity: 0.85; font-size: 13px; }
  .body { padding: 32px 40px; }
  .section { margin-bottom: 24px; }
  .section h2 { font-size: 16px; color: #1a3c5e; border-bottom: 2px solid #2c6496;
                padding-bottom: 6px; margin-bottom: 12px; }
  .content { white-space: pre-wrap; background: #f8f9fa; border-left: 4px solid #2c6496;
             padding: 16px; border-radius: 0 4px 4px 0; font-size: 13px; }
  .badge { display: inline-block; padding: 3px 10px; border-radius: 12px;
           font-size: 11px; font-weight: 600; background: #e8f0fb; color: #2c6496; margin-bottom: 16px; }
  .footer { padding: 16px 40px; background: #f0f0f0; font-size: 12px; color: #888; text-align: center; }
  @media print { body { background: #fff; padding: 0; } .container { box-shadow: none; } }
</style>
</head>
<body>
<div class="container">
  <div class="header">
    <span class="badge">{{ analysis_type|upper }}</span>
    <h1>{{ title }}</h1>
    <div class="meta">Generated: {{ timestamp }} &nbsp;|&nbsp; Vision Assistant</div>
  </div>
  <div class="body">
    <div class="section">
      <h2>Analysis Result</h2>
      <div class="content">{{ content }}</div>
    </div>
    {% if metadata %}
    <div class="section">
      <h2>Metadata</h2>
      {% for key, value in metadata.items() %}
      <p><strong>{{ key }}:</strong> {{ value }}</p>
      {% endfor %}
    </div>
    {% endif %}
  </div>
  <div class="footer">Powered by Vision Assistant &mdash; nvidia/nemotron-3-nano-omni via LM Studio</div>
</div>
</body>
</html>""")


class ReportService:
    def __init__(self):
        _ensure_storage()
        self.storage = Path(settings.storage_path)
        self.reports_dir = self.storage / "reports"

    def _report_path(self, report_id: str, fmt: str) -> Path:
        return self.reports_dir / f"{report_id}.{fmt}"

    def _generate_id(self) -> str:
        return f"rpt_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:6]}"

    # ─── PDF ───────────────────────────────────────────────────────────────────

    def generate_pdf(self, report_id: str, title: str, content: str,
                     analysis_type: str, metadata: dict) -> Path:
        path = self._report_path(report_id, "pdf")
        doc = SimpleDocTemplate(str(path), pagesize=A4,
                                leftMargin=2*cm, rightMargin=2*cm,
                                topMargin=2*cm, bottomMargin=2*cm)
        s = _styles()

        story = []
        story.append(Paragraph(title, s["title"]))
        story.append(Paragraph(f"<font color='grey' size='9'>Type: {analysis_type.upper()} &nbsp;|&nbsp; {datetime.now().strftime('%Y-%m-%d %H:%M')}</font>", s["meta"]))
        story.append(Spacer(1, 0.3*cm))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#2c6496")))
        story.append(Spacer(1, 0.4*cm))

        story.append(Paragraph("Analysis Result", s["heading"]))
        story.append(Paragraph(content.replace("\n", "<br/>"), s["body"]))
        story.append(Spacer(1, 0.3*cm))

        if metadata:
            story.append(Paragraph("Metadata", s["heading"]))
            rows = [[k, str(v)] for k, v in metadata.items()]
            t = Table(rows, colWidths=[5*cm, 11*cm])
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#e8f0fb")),
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.lightgrey),
                ("ROWBACKGROUNDS", (0, 0), (-1, -1), [colors.white, colors.HexColor("#f8f9fa")]),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]))
            story.append(t)

        story.append(Spacer(1, 0.5*cm))
        story.append(Paragraph(
            "<font color='grey' size='8'>Powered by Vision Assistant &mdash; nvidia/nemotron-3-nano-omni via LM Studio</font>",
            s["meta"]
        ))

        doc.build(story)
        return path

    # ─── DOCX ─────────────────────────────────────────────────────────────────

    def generate_docx(self, report_id: str, title: str, content: str,
                      analysis_type: str, metadata: dict) -> Path:
        path = self._report_path(report_id, "docx")
        doc = Document()

        # Title
        h = doc.add_heading(title, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x5E)
            run.font.size = Pt(20)

        # Meta
        p = doc.add_paragraph()
        p.add_run(f"Type: {analysis_type.upper()}  |  ").font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        p.add_run(datetime.now().strftime("%Y-%m-%d %H:%M")).font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        p.runs[0].font.size = Pt(9)
        if p.runs[-1].text:
            p.runs[-1].font.size = Pt(9)

        doc.add_paragraph("_" * 80)

        # Content
        doc.add_heading("Analysis Result", level=2)
        for run in doc.add_paragraph(content).runs:
            run.font.size = Pt(11)

        if metadata:
            doc.add_heading("Metadata", level=2)
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Field"
            hdr[1].text = "Value"
            for cell in hdr:
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                cell._element.get_or_add_tcPr()
                from docx.oxml import OxmlElement
                tc = cell._element
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "2C6496")
                tc.get_or_add_tcPr().append(shd)

            for k, v in metadata.items():
                row = table.add_row().cells
                row[0].text = k
                row[1].text = str(v)

        doc.add_paragraph()
        p = doc.add_paragraph("Powered by Vision Assistant \u2014 nvidia/nemotron-3-nano-omni via LM Studio")
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        doc.save(path)
        return path

    # ─── HTML ──────────────────────────────────────────────────────────────────

    def generate_html(self, report_id: str, title: str, content: str,
                      analysis_type: str, metadata: dict) -> Path:
        path = self._report_path(report_id, "html")
        tmpl = _html_template()
        html = tmpl.render(
            title=title,
            content=content,
            analysis_type=analysis_type,
            timestamp=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            metadata=metadata,
        )
        path.write_text(html, encoding="utf-8")
        return path

    # ─── All formats ───────────────────────────────────────────────────────────

    def generate_all(
        self,
        title: str,
        content: str,
        analysis_type: str = "general",
        metadata: dict | None = None,
    ) -> list[tuple[str, Path]]:
        meta = metadata or {}
        rid = self._generate_id()
        results = []
        for fmt in ["pdf", "docx", "html"]:
            method = getattr(self, f"generate_{fmt}")
            p = method(rid, title, content, analysis_type, meta)
            results.append((fmt, p))
        return results


report_service = ReportService()


# Helper for XML namespace (needed for docx shading)
def qn(tag: str) -> str:
    return f"{{{tag}}}"
